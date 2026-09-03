from typing import Dict, Any
import hashlib
import os
import logging
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

class DocxProcessor:
    """Processor for DOCX documents"""
    
    def __init__(self, config):
        self.config = config
    
    async def process(
        self,
        file_path: str,
        metadata: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """Process a DOCX file and extract text"""
        
        result = {
            "content": "",
            "metadata": metadata or {},
            "word_count": 0,
            "page_count": 0  # Not applicable for DOCX
        }
        
        try:
            # Load the DOCX document
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            result["content"] = "\n".join(full_text)
            result["word_count"] = len(result["content"].split())
            
            # Generate file hash
            with open(file_path, 'rb') as f:
                file_hash = hashlib.md5(f.read()).hexdigest()
                result["file_hash"] = file_hash
            
            return result
            
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            raise
